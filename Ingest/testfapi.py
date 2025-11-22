# main.py
import os
import io
import hashlib
from datetime import datetime, timezone
from typing import Optional, Literal, List, Tuple, Dict, Any

from fastapi import FastAPI, Request, UploadFile, HTTPException, Header, Query, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse

import mimetypes
import json
import base64
import http.client
from urllib.parse import urlparse

# --- cargar .env muy temprano ---
from pathlib import Path
from dotenv import load_dotenv
load_dotenv(dotenv_path=Path(__file__).with_name(".env"), override=True)


APP_VERSION = "upload-robusto-v3"
print("[BOOT]", APP_VERSION)


# =========================
# Helpers Chroma-safe metadata (sin None, solo primitivos o JSON string)
# =========================
PRIMITIVE_TYPES = (str, int, float, bool)

def _to_chroma_primitive(value):
    """Chroma solo acepta str/int/float/bool. None -> '', listas/dicts -> JSON string."""
    if value is None:
        return ""
    if isinstance(value, PRIMITIVE_TYPES):
        return value
    try:
        return json.dumps(value, ensure_ascii=False)
    except Exception:
        return str(value)

def sanitize_metadata_dict(meta: dict) -> dict:
    """Devuelve un dict solo con tipos permitidos por Chroma (sin None)."""
    clean = {}
    for k, v in (meta or {}).items():
        clean[k] = _to_chroma_primitive(v)
    return clean

# =========================
# LangChain splitter
# =========================
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    HAS_LC_SPLITTER = True
except Exception:
    HAS_LC_SPLITTER = False

# =========================
# PDF backends (sin PyMuPDF)
# =========================
# Texto nativo por página: pdfminer.six
try:
    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LTTextContainer
    HAS_PDFMINER = True
except Exception:
    HAS_PDFMINER = False

# Raster por página para OCR total: pypdfium2 (recomendado en Windows/Py3.13)
try:
    import pypdfium2 as pdfium
    HAS_PDFIUM = True
except Exception:
    HAS_PDFIUM = False

# =========================
# Otros formatos
# =========================
# DOCX
try:
    import docx  # python-docx
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# XLSX/CSV
try:
    import pandas as pd
    HAS_PANDAS = True
except Exception:
    HAS_PANDAS = False

# Imágenes (solo para asegurar Pillow instalado)
try:
    from PIL import Image  # noqa: F401
    HAS_PIL = True
except Exception:
    HAS_PIL = False

# =========================
# CONFIGURACIÓN DESDE .ENV
# =========================
OCR_ENDPOINT = os.getenv("OCR_ENDPOINT", "http://127.0.0.1:9999/ocr").strip()
OCR_ENGINE_NAME = os.getenv("OCR_ENGINE_NAME", "tesseract-ocr").strip()
OCR_MODE_DEFAULT: Literal["auto", "force", "off"] = os.getenv("OCR_MODE_DEFAULT", "auto").strip()

CHUNK_SIZE_DEFAULT = int(os.getenv("CHUNK_SIZE", "1500"))
CHUNK_OVERLAP_DEFAULT = int(os.getenv("CHUNK_OVERLAP", "200"))

PDF_TEXT_THRESHOLD_DEFAULT = int(os.getenv("PDF_TEXT_THRESHOLD", "500"))
EMIT_MULTI_DEFAULT = os.getenv("EMIT_MULTI_DEFAULT", "false").strip().lower() == "true"

print(f"[CONFIG] OCR_ENDPOINT = {OCR_ENDPOINT}")
print(f"[CONFIG] OCR_ENGINE_NAME = {OCR_ENGINE_NAME}")
print(f"[CONFIG] OCR_MODE_DEFAULT = {OCR_MODE_DEFAULT}")
print(f"[CONFIG] CHUNK_SIZE_DEFAULT = {CHUNK_SIZE_DEFAULT}")
print(f"[CONFIG] CHUNK_OVERLAP_DEFAULT = {CHUNK_OVERLAP_DEFAULT}")
print(f"[CONFIG] PDF_TEXT_THRESHOLD_DEFAULT = {PDF_TEXT_THRESHOLD_DEFAULT}")
print(f"[CONFIG] EMIT_MULTI_DEFAULT = {EMIT_MULTI_DEFAULT}")

# =========================
# Utilidades
# =========================
def sha256_bytes(b: bytes) -> str:
    h = hashlib.sha256()
    h.update(b)
    return h.hexdigest()

def now_iso_utc() -> str:
    return datetime.now(timezone.utc).isoformat()

def guess_ext_from_mime(content_type: Optional[str]) -> str:
    if not content_type:
        return ".bin"
    ct = content_type.split(";")[0].strip().lower()
    ext = mimetypes.guess_extension(ct)
    manual_map = {
        "application/pdf": ".pdf",
        "text/plain": ".txt",
        "text/csv": ".csv",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
        "image/png": ".png",
        "image/jpeg": ".jpg",
    }
    return manual_map.get(ct, ext or ".bin")

def resolve_filename(x_filename: Optional[str], content_type: Optional[str]) -> str:
    name = (x_filename or "").strip()
    if not name or name.endswith("."):
        name = f"upload{guess_ext_from_mime(content_type)}"
    return os.path.basename(name)

def build_recursive_splitter(chunk_size: int, overlap: int):
    if not HAS_LC_SPLITTER:
        raise HTTPException(500, "Falta langchain_text_splitters para el prechunking")
    return RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", ". ", " ", ""]
    )

def chunk_text(text: str, chunk_size: int, overlap: int) -> List[Tuple[int, int, str]]:
    """Usa LangChain para cortar pero preserva offsets [char_start, char_end)."""
    if not text:
        return []
    splitter = build_recursive_splitter(chunk_size, overlap)
    chunks = splitter.split_text(text)
    res: List[Tuple[int, int, str]] = []
    cursor = 0
    for ch in chunks:
        idx = text.find(ch, cursor)
        if idx == -1:
            idx = text.find(ch)
        if idx == -1:
            idx = cursor
        start = idx
        end = idx + len(ch)
        res.append((start, end, ch))
        cursor = end
    return res

# =========================
# Cliente OCR genérico (llama OCR_ENDPOINT)
# =========================
def call_ocr_endpoint(image_bytes: bytes) -> Tuple[str, Optional[float], Optional[str]]:
    """
    POST {OCR_ENDPOINT} body: {"image_b64": "..."} -> {"text": "...", "confidence": 0..1}
    Devuelve (texto, confianza|None, error|None)
    """
    if not OCR_ENDPOINT:
        return "", None, "no_endpoint"

    parsed = urlparse(OCR_ENDPOINT)
    secure = parsed.scheme == "https"
    conn_class = http.client.HTTPSConnection if secure else http.client.HTTPConnection
    host = parsed.hostname
    port = parsed.port or (443 if secure else 80)
    path = parsed.path or "/"
    if not host:
        return "", None, "endpoint_invalid"

    payload = json.dumps({"image_b64": base64.b64encode(image_bytes).decode("utf-8")})
    headers = {"Content-Type": "application/json"}

    conn = conn_class(host, port, timeout=120)
    try:
        conn.request("POST", path, body=payload, headers=headers)
        resp = conn.getresponse()
        data = resp.read()
        if resp.status != 200:
            return "", None, f"http_{resp.status}"
        jd = json.loads(data.decode("utf-8"))
        text = jd.get("text", "") or ""
        conf = jd.get("confidence", None)
        return text, conf, None
    except Exception as e:
        return "", None, f"exc:{type(e).__name__}"
    finally:
        try:
            conn.close()
        except Exception:
            pass

# =========================
# Extractores por tipo
# =========================
def extract_from_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:
        return content.decode("latin1", errors="ignore")

def extract_from_docx(content: bytes) -> str:
    if not HAS_DOCX:
        raise HTTPException(500, "Falta dependencia python-docx para DOCX")
    f = io.BytesIO(content)
    d = docx.Document(f)
    return "\n".join(p.text for p in d.paragraphs)

def extract_from_xlsx_or_csv(filename: str, content: bytes) -> str:
    if not HAS_PANDAS:
        raise HTTPException(500, "Falta dependencia pandas/openpyxl para XLSX/CSV")
    ext = (filename or "").lower().split(".")[-1]
    bio = io.BytesIO(content)

    # nuevos límites
    CSV_MAX_ROWS = int(os.getenv("CSV_MAX_ROWS", "2000"))
    XLSX_MAX_ROWS = int(os.getenv("XLSX_MAX_ROWS", "2000"))
    XLSX_MAX_SHEETS = int(os.getenv("XLSX_MAX_SHEETS", "3"))

    if ext == "csv":
        # lee en chunks o recorta tras CSV_MAX_ROWS
        df = pd.read_csv(
                      bio,
                    dtype=str,
                    engine="python",
                    nrows=CSV_MAX_ROWS,
                    keep_default_na=False,
                    na_values=[],
                    on_bad_lines="skip",        # ignora filas con columnas extra
                    sep=None                    # autodetecta separador (coma, punto y coma, tab, etc.)
                )

        out = df.to_csv(index=False)
        out = f"=== CSV procesado (máx {CSV_MAX_ROWS} filas, filas mal formateadas omitidas) ===\n" + out
        return out
    else:
        xls = pd.ExcelFile(bio)
        parts = []
        for i, sheet in enumerate(xls.sheet_names[:XLSX_MAX_SHEETS]):
            df = xls.parse(sheet_name=sheet, dtype=str, nrows=XLSX_MAX_ROWS).fillna("")
            chunk = df.to_csv(index=False)
            parts.append(f"=== Hoja: {sheet} (recortada a {len(df)} filas) ===\n{chunk}")
        if len(xls.sheet_names) > XLSX_MAX_SHEETS:
            parts.append(f"=== Aviso: {len(xls.sheet_names)-XLSX_MAX_SHEETS} hojas omitidas por límite ===")
        return "\n\n".join(parts)


def _pdfminer_text_by_pages(content: bytes) -> List[str]:
    """Devuelve lista de textos (uno por página) usando pdfminer.six."""
    if not HAS_PDFMINER:
        raise HTTPException(500, "Falta pdfminer.six para procesar PDF")
    pages_text: List[str] = []
    for layout in extract_pages(io.BytesIO(content)):
        page_parts = []
        for element in layout:
            if isinstance(element, LTTextContainer):
                page_parts.append(element.get_text())
        pages_text.append("".join(page_parts).strip())
    if not pages_text:
        pages_text = [""]
    return pages_text

def _pdfium_raster_page_to_png_bytes(content: bytes, page_index: int, scale: float = 2.0) -> Optional[bytes]:
    """Rasteriza una página específica a PNG usando pypdfium2. Devuelve bytes o None."""
    if not HAS_PDFIUM:
        return None
    pdf = pdfium.PdfDocument(io.BytesIO(content))
    try:
        if page_index < 0 or page_index >= len(pdf):
            return None
        page = pdf[page_index]
        bitmap = page.render(scale=scale).to_pil()
        out = io.BytesIO()
        bitmap.save(out, format="PNG")
        return out.getvalue()
    finally:
        try:
            pdf.close()
        except Exception:
            pass

def extract_pdf_with_policy(content: bytes,
                            ocr_mode: Literal["auto", "force", "off"],
                            pdf_text_threshold: int) -> Tuple[str, Dict[str, Any]]:
    """
    Estrategia:
      - Obtener texto nativo por página con pdfminer.six.
      - Modo auto: si len(texto_nativo) < umbral -> OCR total de la página (pypdfium2);
                   si >= umbral -> solo nativo.
      - Modo force: OCR total de todas las páginas (pypdfium2).
      - Modo off: solo texto nativo.
    """
    pages_meta: List[Dict[str, Any]] = []
    ocr_confs: List[float] = []
    ocr_errors: List[str] = []
    unified_text_parts: List[str] = []

    native_pages = _pdfminer_text_by_pages(content)

    for p_idx, native_text in enumerate(native_pages):
        page_text_parts = []
        page_source = "digital"
        ocr_regions = 0

        do_ocr_total = (ocr_mode == "force") or (ocr_mode == "auto" and len((native_text or "").strip()) < pdf_text_threshold)

        if ocr_mode == "off":
            page_text_parts.append(native_text)
            page_source = "digital"
        elif do_ocr_total:
            png_bytes = _pdfium_raster_page_to_png_bytes(content, p_idx, scale=2.0)
            if png_bytes is not None:
                t, c, err = call_ocr_endpoint(png_bytes)
                if err:
                    ocr_errors.append(f"p{p_idx+1}:{err}")
                if c is not None:
                    try:
                        ocr_confs.append(float(c))
                    except Exception:
                        pass
                page_text_parts.append(t)
                page_source = "ocr"
                ocr_regions = 1
            else:
                # Sin pypdfium2: fallback a texto nativo
                page_text_parts.append(native_text)
                page_source = "digital"
        else:
            # Auto con texto suficiente: solo texto nativo
            page_text_parts.append(native_text)
            page_source = "digital"

        page_text = "\n".join([s for s in page_text_parts if s is not None])
        unified_text_parts.append(f"\n\n=== Página {p_idx+1} ===\n{page_text}")
        pages_meta.append({
            "page_index": p_idx + 1,
            "page_source": page_source,
            "ocr_regions": ocr_regions
        })

    unified_text = "\n".join(unified_text_parts).strip()
    ocr_conf_avg = sum(ocr_confs) / len(ocr_confs) if ocr_confs else None
    used_ocr = any(p.get("page_source") == "ocr" for p in pages_meta)
    page_meta = {
        "pages": pages_meta,
        "ocr_engine": (OCR_ENGINE_NAME if (OCR_ENDPOINT and ocr_mode != "off" and used_ocr) else None),
        "ocr_confidence_avg": ocr_conf_avg if used_ocr else None,
    }
    if ocr_errors and used_ocr:
        page_meta["ocr_errors"] = ocr_errors
    return unified_text, page_meta

def extract_from_image(image_bytes: bytes,
                       ocr_mode: Literal["auto", "force", "off"]) -> Tuple[str, Dict[str, Any]]:
    page_meta = {
        "pages": [{"page_index": 1, "page_source": "ocr" if ocr_mode != "off" else "digital", "ocr_regions": 1 if ocr_mode != "off" else 0}],
        "ocr_engine": OCR_ENGINE_NAME if ocr_mode != "off" else None,
        "ocr_confidence_avg": None
    }
    if ocr_mode == "off":
        return "", page_meta
    t, c, err = call_ocr_endpoint(image_bytes)
    if c is not None:
        page_meta["ocr_confidence_avg"] = c
    if err:
        page_meta["ocr_errors"] = [err]
    return t, page_meta

# =========================
# Helpers OCR limpieza
import re
import unicodedata

def normalize_ocr_text(s: str) -> str:
    if not s:
        return s
    # Unicode y tildes coherentes
    s = unicodedata.normalize("NFC", s)

    # Quitar espacios pegados tras signos/puntuación: "Lenguaje:claro" -> "Lenguaje: claro"
    s = re.sub(r'(:|;|,)(\S)', r'\1 \2', s)

    # Unir líneas duras dentro de párrafos: varias \n -> un salto, o espacio si no hay viñetas
    s = re.sub(r'([^\n])\n(?!\n)', r'\1 ', s)         # salto simple -> espacio
    s = re.sub(r'\n{2,}', '\n\n', s)                  # colapsar múltiples saltos

    # Arreglos típicos de OCR español:
    s = s.replace('¿ ', '¿').replace(' ¡', '¡')       # signos invertidos pegados
    s = re.sub(r'(\S)\*', r'\1*', s)                  # asteriscos sueltos
    s = re.sub(r'\s{2,}', ' ', s)                     # espacios múltiples

    return s.strip()


# =========================
# Trazabilidad (páginas → spans)
# =========================
def build_page_spans(unified_text: str) -> List[Tuple[int, int, int]]:
    """Detecta '=== Página X ===' para mapear (page_no, start_char, end_char)."""
    spans = []
    pos = 0
    current_page = 1
    lines = unified_text.splitlines(keepends=True)
    start_pos_of_current = 0
    for i, ln in enumerate(lines):
        if ln.startswith("=== Página "):
            if i > 0:
                end_pos = pos
                spans.append((current_page, start_pos_of_current, end_pos))
                try:
                    pgstr = ln.strip().split("Página ")[1].split("===")[0].strip()
                    current_page = int(pgstr)
                except Exception:
                    current_page += 1
                start_pos_of_current = pos + len(ln)
        pos += len(ln)
    spans.append((current_page, start_pos_of_current, len(unified_text)))
    return spans

def compute_page_span_for_char_range(char_start: int,
                                     char_end: int,
                                     page_spans: List[Tuple[int, int, int]]) -> Tuple[Optional[int], Optional[int]]:
    touched = [p for (p, s, e) in page_spans if not (char_end <= s or char_start >= e)]
    if not touched:
        return (None, None)
    return (min(touched), max(touched))

# =========================
# Metadatos doc/chunk
# =========================
def build_doc_and_chunk_metadata(filename: str,
                                 content_type: str,
                                 raw: bytes,
                                 unified_text: str,
                                 ocr_mode: str,
                                 chunk_size: int,
                                 chunk_overlap: int,
                                 page_meta: Dict[str, Any],
                                 collection_name: Optional[str],
                                 source: Optional[str],
                                 page_spans: List[Tuple[int, int, int]]
                                 ) -> Tuple[Dict[str, Any], List[Dict[str, Any]], Any]:
    doc_id = sha256_bytes(raw)
    base_meta = {
        "doc_id": doc_id,
        "filename": filename,
        "content_type": content_type or "application/octet-stream",
        "size": len(raw),
        "created_at": now_iso_utc(),
        "ocr_mode": ocr_mode,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "total_chunks": None,
        "collection_name": collection_name,
        "source": source,
        "pages": page_meta.get("pages"),
        "ocr_engine": page_meta.get("ocr_engine"),
        "ocr_confidence_avg": page_meta.get("ocr_confidence_avg"),
        "ocr_errors": page_meta.get("ocr_errors"),
    }

    ocr_pages = {p["page_index"] for p in (page_meta.get("pages") or []) if p.get("page_source") in ("ocr", "mixto")}

    def _build_for_chunk(idx: int, total: int, s: int, e: int) -> Dict[str, Any]:
        p_start, p_end = compute_page_span_for_char_range(s, e, page_spans)
        ocr_used = False
        if p_start is not None and p_end is not None:
            for pn in range(p_start, p_end + 1):
                if pn in ocr_pages:
                    ocr_used = True
                    break
        meta = {
            "chunk_index": idx,
            "total_chunks": total,
            "char_start": s,
            "char_end": e,
            "page_start": p_start,
            "page_end": p_end,
            "ocr_used": ocr_used,
            "doc_id": base_meta["doc_id"],
            "filename": base_meta["filename"],
            "content_type": base_meta["content_type"],
            "size": base_meta["size"],
            "created_at": base_meta["created_at"],
            "collection_name": base_meta["collection_name"],
            "source": base_meta["source"],
            "ocr_mode": base_meta["ocr_mode"],
            "chunk_size": base_meta["chunk_size"],
            "chunk_overlap": base_meta["chunk_overlap"],
        }
        return meta

    return base_meta, [], _build_for_chunk

# =========================
# Pipeline archivo → texto → chunks (LangChain)
# =========================

import re, unicodedata

NORMALIZE_OCR = os.getenv("NORMALIZE_OCR", "true").lower() == "true"

def normalize_ocr_text(s: str) -> str:
    if not s:
        return s
    s = unicodedata.normalize("NFC", s)
    # Espacio tras puntuación pegada por OCR: "Lenguaje:claro,sencillo" -> "Lenguaje: claro, sencillo"
    s = re.sub(r'([:;,])(\S)', r'\1 \2', s)
    # Saltos de línea simples → espacio; preserva párrafos dobles
    s = re.sub(r'([^\n])\n(?!\n)', r'\1 ', s)
    s = re.sub(r'\n{2,}', '\n\n', s)
    # Ajustes típicos ES
    s = s.replace('¿ ', '¿').replace(' ¡', '¡')
    # Colapsa espacios múltiples
    s = re.sub(r'\s{2,}', ' ', s)
    return s.strip()


def process_file_bytes(filename: str,
                       content_type: str,
                       raw: bytes,
                       ocr_mode: Literal["auto", "force", "off"],
                       chunk_size: int,
                       chunk_overlap: int,
                       pdf_text_threshold: int) -> Tuple[str, Dict[str, Any], List[Tuple[int, int, str]], List[Dict[str, Any]]]:
    
    name = (filename or "").lower()
    # Normalizamos el content-type para evitar errores de mayúsculas o parámetros extra
    ctype = (content_type or "").lower().split(";")[0].strip()

    page_meta: Dict[str, Any] = {"pages": [{"page_index": 1, "page_source": "digital", "ocr_regions": 0}], "ocr_engine": None, "ocr_confidence_avg": None}

    # --- CORRECCIÓN AQUÍ: Mirar extensión O Content-Type ---
    if name.endswith(".pdf") or ctype == "application/pdf":
        unified_text, page_meta = extract_pdf_with_policy(raw, ocr_mode, pdf_text_threshold)
        
    elif name.endswith(".docx") or ctype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        unified_text = extract_from_docx(raw)
        
    elif name.endswith((".xlsx", ".csv")) or ctype in ["text/csv", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
        unified_text = extract_from_xlsx_or_csv(filename, raw)
        
    elif name.endswith((".png", ".jpg", ".jpeg")) or ctype.startswith("image/"):
        unified_text, page_meta = extract_from_image(raw, ocr_mode)
        
    else:
        # Fallback a texto plano solo si no es nada de lo anterior
        unified_text = extract_from_txt(raw)

        # 🔹 Normalizar SOLO si es OCR/imagen/PDF (no aplicar a CSV/XLSX)
    if NORMALIZE_OCR:
        is_ocrish = (
            name.endswith((".pdf", ".png", ".jpg", ".jpeg"))
            or bool((page_meta or {}).get("ocr_engine"))  # hubo OCR en PDF auto/force
        )
        if is_ocrish:
            unified_text = normalize_ocr_text(unified_text)

    page_spans = build_page_spans(unified_text)
    chunks = chunk_text(unified_text, chunk_size, chunk_overlap)

    doc_meta, chunk_meta_list, build_chunk_meta = build_doc_and_chunk_metadata(
        filename=filename,
        content_type=content_type,
        raw=raw,
        unified_text=unified_text,
        ocr_mode=ocr_mode,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        page_meta=page_meta,
        collection_name=None,
        source=None,
        page_spans=page_spans
    )

    total = len(chunks)
    doc_meta["total_chunks"] = total
    for idx, (s, e, _) in enumerate(chunks):
        chunk_meta_list.append(build_chunk_meta(idx=idx, total=total, s=s, e=e))

    return unified_text, doc_meta, chunks, chunk_meta_list

# =========================
# FastAPI: endpoints compatibles con OWUI
# =========================
app = FastAPI(title="Ingesta + OCR Selectivo + Prechunking (LangChain) — compatible OWUI")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"]
)

@app.post("/debug/form")
async def debug_form(request: Request):
    ctype = request.headers.get("content-type")
    try:
        form = await request.form()
        items = []
        for k, v in form.multi_items():
            items.append({
                "key": k,
                "type": type(v).__name__,
                "is_upload": hasattr(v, "filename"),
                "filename": getattr(v, "filename", None),
                "content_type": getattr(v, "content_type", None),
                "size_hint": getattr(getattr(v, "spool_max_size", None), "__class__", type("x",(object,),{})()).__name__
            })
        return {"content_type": ctype, "items": items}
    except Exception as e:
        return {"content_type": ctype, "error": repr(e)}


@app.get("/v1/health")
def health():
    return {"status": "ok"}

def build_payload_single(text: str,
                         chunks: List[Tuple[int, int, str]],
                         chunk_meta_list: List[Dict[str, Any]],
                         doc_meta: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "page_content": text,
        "metadata": sanitize_metadata_dict(doc_meta),
        "chunks_meta": [sanitize_metadata_dict(m) for m in chunk_meta_list]
    }

def build_payload_multi(chunks: List[Tuple[int, int, str]],
                        chunk_meta_list: List[Dict[str, Any]]) -> Dict[str, Any]:
    docs = []
    for (s, e, ch_text), meta in zip(chunks, chunk_meta_list):
        docs.append({
            "page_content": ch_text,
            "metadata": sanitize_metadata_dict(meta)
        })
    return {"documents": docs}



async def handle_request(request: Request,
                         x_filename: Optional[str],
                         include_multi: bool,
                         ocr_mode: Literal["auto","force","off"],
                         chunk_size: int,
                         chunk_overlap: int,
                         collection_name: Optional[str],
                         source: Optional[str],
                         pdf_text_threshold: int,
                         upfile_opt: Optional[UploadFile] = None) -> JSONResponse:

    ctype = (request.headers.get("content-type") or "").lower()

    file_bytes = None
    filename = None
    content_type = None

    # 0) Si el endpoint ya nos pasó un UploadFile (de cualquier clase), úsalo directo
    if upfile_opt is not None and hasattr(upfile_opt, "read"):
        file_bytes = await upfile_opt.read()
        if not file_bytes:
            raise HTTPException(400, "Archivo vacío (UploadFile)")
        filename = getattr(upfile_opt, "filename", None) or resolve_filename(
            x_filename, getattr(upfile_opt, "content_type", None)
        )
        content_type = getattr(upfile_opt, "content_type", None) or "application/octet-stream"

    # 1) Si no vino inyectado y el content-type es multipart, intentamos parsear
    elif ctype.startswith("multipart/form-data"):
        form = await request.form()
        upfile = None
        v = form.get("file")
        if hasattr(v, "read"):
            upfile = v
        if upfile is None:
            v = form.get("files")
            if hasattr(v, "read"):
                upfile = v
            else:
                try:
                    for item in form.getlist("files"):
                        if hasattr(item, "read"):
                            upfile = item
                            break
                except Exception:
                    pass
        if upfile is None:
            for _, val in form.multi_items():
                if hasattr(val, "read"):
                    upfile = val
                    break
        if upfile is None:
            raise HTTPException(400, "Se esperaba campo 'file' en multipart/form-data")

        file_bytes = await upfile.read()
        if not file_bytes:
            raise HTTPException(400, "Archivo vacío")
        filename = getattr(upfile, "filename", None) or resolve_filename(
            x_filename, getattr(upfile, "content_type", None)
        )
        content_type = getattr(upfile, "content_type", None) or "application/octet-stream"

    # 2) Cualquier otro content-type: raw bytes
    else:
        file_bytes = await request.body()
        if not file_bytes:
            raise HTTPException(400, "Archivo vacío")
        content_type = request.headers.get("Content-Type", "application/octet-stream")
        filename = resolve_filename(x_filename, content_type)


    # --- procesa (extrae texto, ocr selectivo, chunking) ---
    unified_text, doc_meta, chunks, chunk_meta_list = process_file_bytes(
        filename=filename,
        content_type=content_type,
        raw=file_bytes,
        ocr_mode=ocr_mode,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        pdf_text_threshold=pdf_text_threshold
    )

    # Inyecta collection/source en metadata
    doc_meta["collection_name"] = collection_name
    doc_meta["source"] = source
    for m in chunk_meta_list:
        m["collection_name"] = collection_name
        m["source"] = source

    # Saneamos metadatos antes de responder a OWUI/Chroma
    for key in ("pages", "ocr_errors"):
        if key in doc_meta and not isinstance(doc_meta[key], (str, int, float, bool)):
            try:
                doc_meta[key] = json.dumps(doc_meta[key], ensure_ascii=False)
            except Exception:
                doc_meta.pop(key, None)

    doc_meta = sanitize_metadata_dict(doc_meta)
    chunk_meta_list = [sanitize_metadata_dict(m) for m in chunk_meta_list]

    # --- Construye payload final ---
    payload = build_payload_multi(chunks, chunk_meta_list) if include_multi \
        else build_payload_single(unified_text, chunks, chunk_meta_list, doc_meta)

    return JSONResponse(content=payload, status_code=200)

@app.put("/process")
async def process_alias(
    request: Request,
    file: Optional[UploadFile] = File(None),  # acepta multipart directo
    x_filename: Optional[str] = Header(None, convert_underscores=False),
    ocr: Literal["auto","force","off"] = Query(OCR_MODE_DEFAULT),
    chunk_size: int = Query(CHUNK_SIZE_DEFAULT, ge=1),
    chunk_overlap: int = Query(CHUNK_OVERLAP_DEFAULT, ge=0),
    emit_multi: bool = Query(EMIT_MULTI_DEFAULT),
    collection_name: Optional[str] = Query(None),
    source: Optional[str] = Query(None),
    pdf_text_threshold: int = Query(PDF_TEXT_THRESHOLD_DEFAULT, ge=0),
):
    return await handle_request(
        request=request,
        x_filename=x_filename,
        include_multi=emit_multi,
        ocr_mode=ocr,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        collection_name=collection_name,
        source=source,
        pdf_text_threshold=pdf_text_threshold,
        upfile_opt=file,
    )

from fastapi import UploadFile  # ya lo tienes
from starlette.datastructures import UploadFile as StarletteUploadFile  # para isinstance robusto

from fastapi import UploadFile  # ya lo tienes
from starlette.datastructures import UploadFile as StarletteUploadFile  # para isinstance robusto

@app.post("/upload")
async def upload_post(
    request: Request,
    x_filename: Optional[str] = Header(None, convert_underscores=False),
    ocr: Literal["auto","force","off"] = Query(OCR_MODE_DEFAULT),
    chunk_size: int = Query(CHUNK_SIZE_DEFAULT, ge=1),
    chunk_overlap: int = Query(CHUNK_OVERLAP_DEFAULT, ge=0),
    emit_multi: bool = Query(EMIT_MULTI_DEFAULT),
    collection_name: Optional[str] = Query(None),
    source: Optional[str] = Query(None),
    pdf_text_threshold: int = Query(PDF_TEXT_THRESHOLD_DEFAULT, ge=0),
):
    # 1) Detección de multipart y extracción del PRIMER UploadFile que encontremos, sin importar la clave.
    ctype = (request.headers.get("content-type") or "").lower()
    if not ctype.startswith("multipart/form-data"):
        raise HTTPException(400, "Este endpoint requiere multipart/form-data con un archivo")

    form = await request.form()

    upfile: Optional[UploadFile] = None
    # a) clave estándar "file"
    v = form.get("file")
    if isinstance(v, UploadFile) or isinstance(v, StarletteUploadFile):
        upfile = v

    # b) clave alternativa "files"
    if upfile is None:
        v = form.get("files")
        if isinstance(v, UploadFile) or isinstance(v, StarletteUploadFile):
            upfile = v
        else:
            try:
                for item in form.getlist("files"):
                    if isinstance(item, UploadFile) or isinstance(item, StarletteUploadFile):
                        upfile = item
                        break
            except Exception:
                pass

    # c) última chance: primer UploadFile en cualquier clave
    if upfile is None:
        for _, val in form.multi_items():
            if isinstance(val, UploadFile) or isinstance(val, StarletteUploadFile):
                upfile = val
                break

    if not isinstance(upfile, (UploadFile, StarletteUploadFile)):
        raise HTTPException(400, "Se esperaba al menos un archivo en multipart/form-data")

    # 2) Pasamos el UploadFile directo al pipeline común
    return await handle_request(
        request=request,
        x_filename=x_filename,
        include_multi=emit_multi,
        ocr_mode=ocr,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        collection_name=collection_name,
        source=source,
        pdf_text_threshold=pdf_text_threshold,
        upfile_opt=upfile,
    )



@app.put("/v1/process")
async def process_v1(
    request: Request,
    file: Optional[UploadFile] = File(None),  # acepta multipart directo
    x_filename: Optional[str] = Header(None, convert_underscores=False),
    ocr: Literal["auto","force","off"] = Query(OCR_MODE_DEFAULT),
    chunk_size: int = Query(CHUNK_SIZE_DEFAULT, ge=1),
    chunk_overlap: int = Query(CHUNK_OVERLAP_DEFAULT, ge=0),
    emit_multi: bool = Query(EMIT_MULTI_DEFAULT),
    collection_name: Optional[str] = Query(None),
    source: Optional[str] = Query(None),
    pdf_text_threshold: int = Query(PDF_TEXT_THRESHOLD_DEFAULT, ge=0),
):
    return await handle_request(
        request=request,
        x_filename=x_filename,
        include_multi=emit_multi,
        ocr_mode=ocr,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        collection_name=collection_name,
        source=source,
        pdf_text_threshold=pdf_text_threshold,
        upfile_opt=file,
    )
